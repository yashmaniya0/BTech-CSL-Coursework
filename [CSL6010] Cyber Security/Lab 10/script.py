from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from PIL import Image
import numpy as np
from tqdm import tqdm
import re


def convert_doc_to_text(path):
    doc = Document(path) # reading the doc file using python-docx
    text = ""
    for para in doc.paragraphs:
        for run in para.runs:
            text += run.text # adding all letters in the text into string
    return text

def hide_image_in_rtf(image_path, rtf_path, output_path):
    with open(image_path, 'rb') as f:
        image_bytes = f.read()

    # Convert image bytes to semagram format
    semagram = ""
    for byte in image_bytes:
        semagram += chr(byte)

    # Load the RTF file contents
    with open(rtf_path, 'r') as f:
        rtf_contents = f.read()

    # Replace the semagram placeholder with the image semagram
    rtf_contents = re.sub(r"\\semagram\s*\\", "\\semagram " + semagram + "\\", rtf_contents)

    # Write the modified RTF contents to the output file
    with open(output_path, 'w') as f:
        f.write(rtf_contents)

def extract_image_from_rtf(rtf_path, output_path):
    # Load the RTF file contents
    with open(rtf_path, 'r') as f:
        rtf_contents = f.read()

    # Extract the semagram from the RTF contents
    semagram = re.search(r"\\semagram\s*(.*?)\\", rtf_contents).group(1)

    # Convert semagram to image bytes
    image_bytes = bytearray()
    for char in semagram:
        image_bytes.append(ord(char))

    # Save image bytes to file
    with open(output_path, 'wb') as f:
        f.write(image_bytes)

def hide_image_in_docx(img,text):
    width, height = img.size
    size = len(text)
    document = Document() # creating new document
    para = document.add_paragraph() # adding paragrah

    # iteration over all the pixels of image
    for col in tqdm(range(height)):
        for row in range(width):
            index = col*width + row # index for each pixel
            word = text[index%size] # find letter corresponding index from the text file
            char = para.add_run(word) # add leter into paragraph
            char.font.size = Pt(1)
            r,g,b = img.getpixel((row,col)) # pixel for each row and column
            char.font.color.rgb = RGBColor(r,g,b) # add color to the letter corresponding pixel of the image

        para.add_run("\n")

    document.save("output.docx") # saving the document

    print(f'Success : Image hidden in output document!!')

img = Image.open("./bald_eagle.jpg")
text = convert_doc_to_text("./text.docx")
hide_image_in_docx(img,text)
